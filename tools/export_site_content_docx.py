#!/usr/bin/env python3
"""Export full Russian site copy to Word for editorial review.

Includes proposed slogan alignments from the live original site (vlasta-s.com).
Changed / restored slogans are shown in green bold.
"""

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SITE = json.loads((ROOT / "content/site.json").read_text(encoding="utf-8"))
CASES = json.loads((ROOT / "content/cases.json").read_text(encoding="utf-8"))
NEWS = json.loads((ROOT / "content/news.json").read_text(encoding="utf-8"))
O = SITE["org"]
FOUNDED = O["founded"]

GREEN = RGBColor(0, 128, 0)
BLUE = RGBColor(0, 102, 204)
GRAY = RGBColor(120, 120, 120)

# ---------------------------------------------------------------------------
# Original slogans / statements from https://vlasta-s.com (Aug 2026 snapshot)
# Only include texts that exist verbatim (or nearly) on the live original site.
# ---------------------------------------------------------------------------
ORIGINAL = {
    "heroTitle": "Безопасность бизнеса в надёжных руках",
    "missionSlogan": (
        "Используя передовые технологии, опираясь на высокие моральные ценности "
        "и постоянно стремясь к успеху, обеспечивать экономическую безопасность "
        "бизнеса наших клиентов."
    ),
    "metaDescription": (
        "Защита бренда, корпоративная безопасность, личная безопасность - "
        "Власта консалтинг"
    ),
    "servicesHeading": "Наши услуги",
    "integratedSecurityTitle": "Комплексная безопасность",
    "integratedSecurityConcept": "Концепция комплексной безопасности",
    "integratedSecurityText": (
        "Росту успешности компании всегда сопутствует рост рисков. Они "
        "подразделяются на внешние (воровство как продукции, так и объектов "
        "интеллектуальной собственности) и внутренние (финансовые и иные "
        "махинации, шантаж руководства, физические угрозы персоналу как "
        "следствие недостатка эффективности систем внутренней безопасности). "
        "Наиболее защищенная от рисков компания получает абсолютное "
        "преимущество перед конкурентами, что напрямую влияет на ее финансовые "
        "показатели. При этом чем выше компания находится на уровне своего "
        "развития, тем существеннее для нее потенциальный урон и тем больше "
        "ресурсов затрачивается на их устранение. Наиболее эффективным "
        "вариантом является не реагирование на уже возникшие риски, а создание "
        "комплексной системы обеспечения безопасности бизнеса."
    ),
    "serviceIp": "Интеллектуальная собственность",
    "serviceCorporate": "Корпоративная безопасность",
    "servicePersonal": "Личная безопасность",
    "brandProtectionIntro": (
        "ООО «Власта-Консалтинг» имеет многолетний опыт работы в области "
        "защиты прав на объекты интеллектуальной собственности. В России мы "
        "были практически первыми, кто разработал и запустил в действие в "
        "2006 году программу защиты товарных знаков."
    ),
    "brandProtectionFullCycle": (
        "Защита товарных знаков заключается в планомерной и полномасштабной "
        "работе как по всем возможным направлениям, так и в отношении всех "
        "видов нарушений (тождество, сходство до степени смешения, "
        "недобросовестная конкуренция). Комплексное использование всех "
        "доступных инструментов позволяет обеспечить максимально возможный "
        "уровень защиты товарного знака."
    ),
    "trademarkRegistrationLead": (
        "Товарный знак - это визитная карточка и ценный актив Вашего бизнеса. "
        "Именно товарный знак будет ассоциироваться у клиентов с Вашими "
        "товарами и услугами. При помощи товарного знака можно эффективно "
        "защитить бизнес от подделок, а также получить прибыль от продажи "
        "лицензий третьим лицам. Прежде чем начать использовать товарный знак, "
        "необходимо обеспечить его правовую охрану."
    ),
    "corporateSecurityIntro": (
        "Холдинговая компания Власта Консалтинг предлагает комплексные "
        "решения в области экономической безопасности. Двадцать лет опыта "
        "сделали из нас первоклассную команду, которой наши клиенты готовы "
        "доверить решение любых вопросов в сфере защиты бизнеса."
    ),
    "personalSecurityIntro": (
        "Реалии современного мира диктуют необходимость в услугах по "
        "обеспечению личной безопасности. Комплекс мер, направленных на "
        "сохранение жизни и защиты бизнес инфраструктуры, состоит из множества "
        "элементов, эффективность которых зависит от профессионализма команды."
    ),
    "oneWindowMission": (
        "Общая миссия нашей компании, заключающаяся в обеспечении комплексной "
        "защиты от рисков (формат «одного окна»), находит свое отражение и в "
        "рамках услуг по консультировании по вопросам обеспечения безопасности."
    ),
    "intelligenceIntro": (
        "Наша команда высококвалифицированных профессионалов предоставляет "
        "услуги, объединяющие в себе глубокую аналитику и возможность получать "
        "информацию из проверенных источников в странах бывшего Восточного "
        "блока и за его пределами, чтобы помочь нашим клиентам выстроить "
        "бизнес-стратегии, минимизирующие риски, и закрепить свои позиции в "
        "этом регионе."
    ),
    "whoDesc": (
        "Основанная в 2006 году компания Власта-Консалтинг занимает "
        "лидирующее место в сфере обеспечения безопасности бизнеса и широко "
        "известна в отечественных и иностранных бизнес-кругах. Главный принцип "
        "в работе «Власта-Консалтинг» - комплексный подход к каждому клиенту."
    ),
    "aboutLead": (
        "Многолетний опыт работы позволяет нашим специалистам анализировать "
        "риски, предвидеть неблагоприятные сценарии развития событий и "
        "эффективно реагировать на критические ситуации. Мы работаем и "
        "оказываем услуги как на территории России, так и в странах СНГ, "
        "Европы, Северной и Южной Америки, Африки, Юго-Восточной Азии. "
        "Сотрудничество с зарубежными партнерами наряду с членством в "
        "профильных международных ассоциациях позволяет нам эффективно "
        "представлять в России интересы транснациональных корпораций."
    ),
    "valuesPrinciple": (
        "В своей работе мы руководствуемся принципами законности и правилами "
        "деловой этики, а также осуществляем свою деятельность с максимальной "
        "эффективностью для наших клиентов."
    ),
    "kpiHeader": "Ключевые показатели за год",
    "kpiDamage": "Предотвращенный ущерб",
    "kpiCases": "Уголовных дел",
    "kpiProductions": "Производств закрыто",
    "kpiBudget": "Затрат на программу возвращено правообладателю через ГИ",
    "cookieText": (
        "Продолжая просмотр настоящего сайта, Вы соглашаетесь с использованием "
        "файлов Cookie и иных методов, средств и инструментов интернет-статистики "
        "и настройки, применяемых на сайте для повышения удобства использования "
        "сайта."
    ),
    "hoursLabel": "Будни: 9:30 - 18:00",
    "address": "119048, город Москва, ул. Усачёва, д. 13, помещ. 4н",
}

# Proposed rewrites: key -> text from original (only where applicable)
REWRITES = {
    "intro": ORIGINAL["missionSlogan"],
    "footerAbout": ORIGINAL["missionSlogan"],
    "homeDesc": ORIGINAL["metaDescription"],
    "svcTitle": ORIGINAL["servicesHeading"],
    "integratedSecurityText": ORIGINAL["integratedSecurityText"],
    "integratedSecurityTitle": ORIGINAL["integratedSecurityTitle"],
    "brandIntro": ORIGINAL["brandProtectionIntro"],
    "brandTagline": ORIGINAL["brandProtectionFullCycle"],
    "trademarkRegistrationText": ORIGINAL["trademarkRegistrationLead"],
    "securityIntro": ORIGINAL["corporateSecurityIntro"],
    "personalSecurityIntro": ORIGINAL["personalSecurityIntro"],
    "intelligenceIntro": ORIGINAL["intelligenceIntro"],
    "whoDesc": ORIGINAL["whoDesc"],
    "aboutLead": ORIGINAL["aboutLead"],
    "valuesPrinciple": ORIGINAL["valuesPrinciple"],
    "apprKpiHeader": ORIGINAL["kpiHeader"],
    "cookieText": ORIGINAL["cookieText"],
    "hours": ORIGINAL["hoursLabel"],
    # CTA: no equivalent block on vlasta-s.com — intentionally omitted
}

# Page / UI copy (mirrors tools/build.mjs, ru locale)
T = {
    "nav": "Главная · Услуги · О компании · Кейсы · Новости",
    "contact": "Связаться",
    "contactBtn": "Контакты",
    "menu": "Меню",
    "more": "Подробнее",
    "readOn": "Читать",
    "allServices": "Все услуги",
    "allCases": "Все кейсы",
    "allNews": "Все новости",
    "caseStudy": "Разбор кейса",
    "home": "Главная",
    "nav_services": "Услуги",
    "nav_contacts": "Контакты",
    "nav_cases": "Кейсы",
    "navigation": "Основная навигация",
    "documents": "Документы",
    "info": "Информация",
    "telFax": "Тел./факс",
    "themeLabel": "Переключить тёмную тему",
    "openMenu": "Открыть меню",
    "closeMenu": "Закрыть меню",
    "toTop": "Наверх",
    "crumbs": "Хлебные крошки",
    "rights": "Все права защищены.",
    "privacy": "Политика конфиденциальности",
    "footerAbout": (
        f"Обеспечиваем экономическую безопасность бизнеса и защиту брендов "
        f"от контрафакта с {FOUNDED} года — на передовых технологиях и высоких "
        f"моральных ценностях."
    ),
    "ctaKicker": "Начнём сотрудничество",
    "ctaTitle": "Обсудим, как защитить ваш бизнес",
    "ctaText": (
        "Проведём конфиденциальную консультацию, оценим риски и предложим "
        "решение под вашу задачу."
    ),
    "cookieTitle": "Файлы cookie",
    "cookieText": (
        "Мы используем cookie, чтобы сайт работал корректно и чтобы понимать, "
        "какие материалы вам полезны. Аналитику можно отключить — на работу "
        "сайта это не повлияет."
    ),
    "cookieMore": "Подробнее",
    "ckReject": "Отклонить всё",
    "ckNeeded": "Только необходимые",
    "ckAll": "Принять всё",
    "scrollHint": "листайте",
    "sources": "Источники и упоминания",
    "topics": "Темы кейса",
    "newer": "Новее",
    "earlier": "Ранее",
    "related": "По теме",
    "showMore": "Показать ещё",
    "emptyNews": "По выбранной теме материалов пока нет.",
    "emptyCases": "По выбранной категории кейсов пока нет.",
    "direction": "Направление",
    "region": "Регион",
    "outcome": "Результат",
    "format": "Формат",
    "projectWork": "Проектная работа",
    "regionValue": "Россия · СНГ · ЕАЭС",
    "phone": "Телефон",
    "email": "Электронная почта",
    "address": "Адрес",
    "hours": "Часы работы",
    "callUs": "Позвонить",
    "writeUs": "Написать письмо",
    "weekdays": "Понедельник – Пятница",
    "otherCases": "Другие кейсы",
    "otherNews": "Другие материалы",
}

C = {
    "homeTitle": (
        f"{O['name']} — защита брендов и экономическая безопасность бизнеса"
    ),
    "homeDesc": (
        f"«{O['name']}» с {FOUNDED} года защищает товарные знаки от контрафакта "
        f"и обеспечивает экономическую безопасность бизнеса в России, СНГ и "
        f"странах ЕАЭС: ТРОИС, рейды, проверки контрагентов, сопровождение в суде."
    ),
    "homeKw": (
        "защита бренда, борьба с контрафактом, ТРОИС, экономическая безопасность "
        "бизнеса, проверка контрагентов, бизнес-разведка, Власта-Консалтинг"
    ),
    "heroPill": f"С {FOUNDED} года · Москва · Россия и ЕАЭС",
    "heroTitle": "Безопасность бизнеса\nв надёжных руках",
    "heroScroll": "Пролистать к описанию",
    "intro": (
        "Защищаем бренды от контрафакта и обеспечиваем экономическую "
        "безопасность компаний в России, СНГ и странах ЕАЭС — опираясь на "
        "передовые технологии и высокие моральные ценности."
    ),
    "svcKicker": "Услуги",
    "svcTitle": "Наши направления",
    "svcDesc": (
        "Единая методология — от анализа рисков до сопровождения «под ключ» "
        "в суде. Каждое направление работает самостоятельно и усиливает остальные."
    ),
    "apprKicker": "Подход и практика",
    "apprTitle": "Все отделы — одна система",
    "apprDesc": (
        "Профильные отделы работают в одном контуре — от аналитики и полевых "
        "мероприятий до права и цифровой среды. Координация между ними "
        "обеспечивает правообладателям измеримый эффект наших программ защиты "
        "брендов."
    ),
    "casesKicker": "Кейсы",
    "casesTitle": "Как мы решаем задачи клиентов",
    "histKicker": "История компании",
    "histTitle": "Путь, отмеченный международным признанием",
    "histHint": "От московского старта до международной практики",
    "newsKicker": "Новости",
    "newsTitle": "Компания в публичном пространстве",
    "marquee": "Нам доверяют ведущие российские и международные бренды",
    "svcPageTitle": (
        "Услуги: защита бренда, ТРОИС, проверки контрагентов, "
        "бизнес-разведка — Власта-Консалтинг"
    ),
    "svcPageDesc": (
        "Восемь направлений в четырёх блоках: разведка и анализ, безопасность "
        "бизнеса, защита бренда и ИС, консалтинг. Регистрация в ТРОИС, рейды с "
        "полицией и таможней, комплаенс, сопровождение в суде."
    ),
    "svcH1": "Услуги по защите бренда и безопасности бизнеса",
    "svcLead": (
        "Единая методология: анализ рисков, предупреждение угроз и сопровождение "
        "клиента вплоть до защиты интересов в суде — в России, СНГ и странах ЕАЭС."
    ),
    "aboutH1": f"Эксперты по защите брендов и безопасности бизнеса с {FOUNDED} года",
    "aboutLead": (
        "Мы помогаем компаниям расти спокойно — анализируем риски, предвидим "
        "неблагоприятные сценарии и выстраиваем системы защиты, которые работают "
        "на опережение."
    ),
    "whoKicker": "Кто мы",
    "whoTitle": "Надёжный партнёр в вопросах экономической безопасности",
    "whoDesc": (
        "Основанная в 2006 году, компания занимает лидирующее место в сфере "
        "обеспечения безопасности бизнеса и защиты интеллектуальной собственности "
        "в России."
    ),
    "teamKicker": "Руководство",
    "teamTitle": "Команда, которая отвечает за результат",
    "assocKicker": "Партнёрство",
    "assocTitle": "Ассоциации и профессиональные сообщества",
    "assocDesc": (
        "Мы состоим в ведущих российских и международных объединениях. Нажмите "
        "на карточку, чтобы узнать об участии в каждой ассоциации."
    ),
    "clientsKicker": "Клиенты",
    "clientsTitle": "Нам доверяют ведущие бренды",
    "clientsNote": "и ещё более 80 брендов под нашей защитой",
    "lettersKicker": "Отзывы",
    "lettersTitle": "Благодарственные письма",
    "lettersDesc": "Нажмите на письмо, чтобы открыть его целиком.",
    "casesH1": "Кейсы: контрафакт, расследования и защита активов",
    "casesLead": (
        "Реальные проекты по защите товарных знаков, антиконтрафактным программам, "
        "проверкам контрагентов и внутренним расследованиям. Детали обезличены в "
        "целях конфиденциальности клиентов."
    ),
    "newsH1": "Новости борьбы с контрафактом и защиты брендов",
    "newsLead": (
        "Антиконтрафактные операции и изъятия, инициативы на площадках ЕЭК и ФТС, "
        "участие в международных форумах по защите интеллектуальной собственности."
    ),
    "contactsKick": "Свяжитесь с нами",
    "contactsH1": "Обсудим безопасность вашего бизнеса",
    "contactsLead": (
        "Проведём конфиденциальную консультацию и предложим решение под вашу задачу."
    ),
    "privacyKick": "Правовая информация",
    "privacyH1": "Политика конфиденциальности",
    "privacyTitle": "Политика конфиденциальности — Власта-Консалтинг",
    "privacyDesc": (
        "Политика обработки персональных данных ООО «Власта-Консалтинг»: какие "
        "данные мы собираем, цели и правовые основания обработки, сроки хранения "
        "и ваши права."
    ),
}

PRIVACY = [
    (
        "Общие положения",
        [
            f"Настоящая Политика определяет порядок обработки персональных данных "
            f"{O['legal']} (далее — Компания) и меры по обеспечению их безопасности.",
            "Используя сайт и обращаясь к нам по телефону или электронной почте, "
            "вы соглашаетесь с условиями настоящей Политики.",
        ],
    ),
    (
        "Какие данные мы обрабатываем",
        [
            "Имя, название компании, телефон, адрес электронной почты и содержание "
            "обращения — в объёме, который вы сообщаете нам сами.",
            "Технические данные: IP-адрес, тип браузера и устройства, источник "
            "перехода, действия на сайте — в обезличенном виде для статистики.",
        ],
    ),
    (
        "Цели обработки",
        [
            "Ответ на ваше обращение и проведение консультации.",
            "Улучшение работы сайта и качества услуг.",
            "Исполнение требований законодательства Российской Федерации.",
        ],
    ),
    (
        "Правовые основания",
        [
            "Обработка осуществляется на основании вашего согласия, а также в "
            "соответствии с Федеральным законом от 27.07.2006 № 152-ФЗ «О "
            "персональных данных».",
        ],
    ),
    (
        "Передача третьим лицам",
        [
            "Компания не продаёт и не передаёт персональные данные третьим лицам, "
            "за исключением случаев, прямо предусмотренных законодательством, либо "
            "когда это необходимо для исполнения вашего обращения.",
        ],
    ),
    (
        "Срок хранения",
        [
            "Персональные данные хранятся не дольше, чем это необходимо для целей "
            "обработки, либо до отзыва вашего согласия.",
        ],
    ),
    (
        "Файлы cookie",
        [
            "Сайт использует файлы cookie для корректной работы интерфейса и сбора "
            "обезличенной статистики. Вы можете отключить cookie в настройках "
            "браузера.",
        ],
    ),
    (
        "Ваши права",
        [
            "Вы вправе запросить сведения об обработке ваших данных, потребовать "
            "их уточнения, блокирования или удаления, а также отозвать согласие на "
            "обработку.",
            f"Для реализации прав направьте обращение на {O['email']}.",
        ],
    ),
    (
        "Контакты",
        [
            f"{O['legal']}, {O['address']}. Телефон: {O['phone']}. "
            f"E-mail: {O['email']}.",
        ],
    ),
]


def norm(s):
    return re.sub(r"\s+", " ", strip_html(s or "")).strip().lower()


def strip_html(text):
    if not text:
        return ""
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"</?b>", "", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    return text.strip()


def add_run(p, text, *, bold=False, color=None, size=11):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return r


def add_plain(doc, text, size=11):
    if not text:
        return
    for para in text.split("\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph()
            add_run(p, para, size=size)


def add_green_bold(doc, text, size=11):
    if not text:
        return
    for para in text.split("\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph()
            add_run(p, para, bold=True, color=GREEN, size=size)


def add_blue(doc, text, size=11):
    if not text:
        return
    for para in text.split("\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph()
            add_run(p, para, bold=True, color=BLUE, size=size)


def add_ab_compare(doc, label, original, current, *, note=None):
    """A/B block: green original (vlasta-s.com), then blue current new-site text."""
    add_h4(doc, label)
    p_a = doc.add_paragraph()
    add_run(p_a, "A — Оригинал (vlasta-s.com):", bold=True, color=GREEN, size=10)
    add_green_bold(doc, strip_html(original) if original else "—")
    p_b = doc.add_paragraph()
    add_run(p_b, "B — Сейчас на новом сайте:", bold=True, color=BLUE, size=10)
    if current and strip_html(current):
        add_blue(doc, strip_html(current))
    else:
        add_blue(doc, "Нет прямого аналога на новом сайте.")
    if note:
        p_n = doc.add_paragraph()
        add_run(p_n, note, color=GRAY, size=9)
    doc.add_paragraph()


def add_labeled_field(doc, label, current, rewrite_key=None):
    """Show A/B comparison when a vlasta-s.com original exists; else plain text."""
    proposed = REWRITES.get(rewrite_key) if rewrite_key else None
    if proposed and norm(current) != norm(proposed):
        add_ab_compare(doc, label, proposed, current)
    elif proposed and norm(current) == norm(proposed):
        add_h4(doc, label)
        add_plain(doc, current)
        p = doc.add_paragraph()
        add_run(
            p,
            "(Совпадает с оригиналом vlasta-s.com — сравнение не требуется.)",
            color=GRAY,
            size=9,
        )
    else:
        add_h4(doc, label)
        add_plain(doc, current)


def add_kicker(doc, text):
    p = doc.add_paragraph(text.upper())
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = RGBColor(107, 110, 134)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_h4(doc, text):
    doc.add_heading(text, level=4)


def add_bullets(doc, items):
    for item in items:
        if item:
            doc.add_paragraph(strip_html(str(item)), style="List Bullet")


def add_section_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def export_preface(doc):
    title = doc.add_heading(
        "Власта-Консалтинг — полный текст сайта для рецензии", level=0
    )
    title.runs[0].font.size = Pt(26)

    add_plain(
        doc,
        "Документ содержит все существенные тексты русской версии нового сайта "
        "(adva.media/vlasta): шапка, подвал, все разделы, кейсы, новости, "
        "политика конфиденциальности и служебные подписи.",
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run(p, "ВАЖНО. ", bold=True)
    add_run(
        p,
        "Сверка с оригинальным сайтом ",
    )
    add_run(p, "vlasta-s.com", bold=True)
    add_run(
        p,
        ". Формат сравнения A/B: сначала ",
    )
    add_run(p, "A — оригинал (зелёный, полужирный)", bold=True, color=GREEN)
    add_run(p, ", сразу под ним ", )
    add_run(p, "B — текущая версия нового сайта (синий, полужирный)", bold=True, color=BLUE)
    add_run(
        p,
        ". Так удобнее править новый текст, глядя на оригинал сверху.",
    )
    doc.add_paragraph()

    add_h3(doc, "Сводка A/B: оригинал vs новый сайт")
    values_current = "\n".join(
        f"{v.get('num', '')}. {v.get('title', '')}: {v.get('text', '')}"
        for v in SITE.get("values", [])
    )
    kpi_current = "\n".join(
        f"{(r.get('prefix', '') or '')}{r.get('value', '')}{r.get('unit', '')} — "
        f"{strip_html(r.get('label', ''))}"
        for r in SITE.get("results", [])
    )
    integrated_dir = next(
        (
            d
            for s in SITE["services"]
            if s["id"] == "security"
            for d in s.get("directions", [])
            if d["title"] == "Комплексная система безопасности"
        ),
        None,
    )
    for label, original, current in [
        ("Главный заголовок (hero)", ORIGINAL["heroTitle"], C["heroTitle"].replace("\n", " ")),
        ("Миссия / слоган (intro + footer)", ORIGINAL["missionSlogan"], C["intro"]),
        ("Слоган в подвале (текущий footer)", ORIGINAL["missionSlogan"], T["footerAbout"]),
        ("Meta description", ORIGINAL["metaDescription"], C["homeDesc"]),
        ("Заголовок блока услуг", ORIGINAL["servicesHeading"], C["svcTitle"]),
        (
            "Комплексная безопасность — заголовок",
            ORIGINAL["integratedSecurityTitle"],
            (integrated_dir or {}).get("title", "Комплексная система безопасности"),
        ),
        (
            "Комплексная безопасность — текст",
            ORIGINAL["integratedSecurityText"],
            (integrated_dir or {}).get("text", ""),
        ),
        ("О компании — абзац 1 (кто мы)", ORIGINAL["whoDesc"], C["whoDesc"]),
        ("О компании — абзац 2 (опыт / география)", ORIGINAL["aboutLead"], C["aboutLead"]),
        (
            "Принципы работы",
            ORIGINAL["valuesPrinciple"],
            values_current,
        ),
        (
            "KPI — заголовок и состав",
            ORIGINAL["kpiHeader"]
            + "\n• "
            + ORIGINAL["kpiDamage"]
            + "\n• "
            + ORIGINAL["kpiCases"]
            + "\n• "
            + ORIGINAL["kpiProductions"]
            + "\n• "
            + ORIGINAL["kpiBudget"],
            "Показатели эффективности (KPI)\n" + kpi_current,
        ),
        ("Cookie", ORIGINAL["cookieText"], T["cookieText"]),
        ("Часы работы", ORIGINAL["hoursLabel"], O["hours"]),
        ("Адрес", ORIGINAL["address"], O["address"]),
    ]:
        add_ab_compare(doc, label, original, current)

    add_h4(doc, "CTA-блок (призыв к действию)")
    add_plain(
        doc,
        "На vlasta-s.com отдельного CTA-баннера нет. Ниже только текущий текст "
        "нового сайта (синим) — без зелёного оригинала.",
    )
    add_blue(doc, f"{T['ctaKicker']}\n{T['ctaTitle']}\n{T['ctaText']}")


def export_chrome(doc):
    add_h2(doc, "Шапка, навигация, подвал, CTA, cookie")
    add_h3(doc, "Верхняя строка (topbar)")
    add_plain(doc, f"Телефон: {O['phone']}")
    add_plain(doc, f"E-mail: {O['email']}")
    add_plain(doc, f"Часы работы: {O['hours']}")
    add_plain(doc, "Переключатель языка: RU / EN")
    add_plain(doc, f"Кнопка: {T['themeLabel']}")

    add_h3(doc, "Навигация")
    add_plain(doc, T["nav"])
    add_plain(doc, f"Кнопка «{T['contactBtn']}» → contacts.html")

    add_h3(doc, "Подвал (footer)")
    add_plain(doc, f"Название: {O['name']} / {O['legal']}")
    add_plain(doc, f"Подпись бренда (tagline): {O['tagline']}")
    add_labeled_field(doc, "О компании (слоган в подвале)", T["footerAbout"], "footerAbout")

    add_h4(doc, "Колонка «Услуги»")
    for s in SITE["services"]:
        add_plain(doc, f"• {s['title']}")

    add_h4(doc, "Колонка «Документы»")
    add_plain(doc, T["privacy"])

    add_h4(doc, "Колонка «Кейсы» (первые 9 в подвале)")
    for c in CASES[:9]:
        add_plain(doc, f"• {c['title']}")

    add_h4(doc, "Колонка «Информация»")
    add_plain(doc, f"{T['hours']}: {O['hours']}")
    add_plain(doc, f"{T['address']}: {O['address']}")
    add_plain(doc, f"{T['telFax']}: {O['phone']}")
    add_plain(doc, f"{T['email']}: {O['email']}")

    add_h4(doc, "Нижняя строка подвала")
    add_plain(doc, f"{T['rights']} © {O['legal']}, [год]")

    add_h3(doc, "Мобильное меню")
    add_plain(doc, T["menu"])
    add_plain(doc, f"{T['openMenu']} / {T['closeMenu']}")

    add_h3(doc, "Блок призыва к действию (CTA на всех страницах)")
    add_plain(
        doc,
        "На vlasta-s.com нет аналога этого блока — зелёная правка не предлагается.",
    )
    add_plain(doc, f"Сейчас: {T['ctaKicker']}")
    add_plain(doc, f"Сейчас: {T['ctaTitle']}")
    add_plain(doc, f"Сейчас: {T['ctaText']}")
    add_plain(doc, f"{T['phone']}: {O['phone']}")
    add_plain(doc, f"{T['email']}: {O['email']}")
    add_plain(doc, f"Кнопка: {T['contactBtn']}")

    add_h3(doc, "Баннер cookie")
    add_plain(doc, T["cookieTitle"])
    add_labeled_field(doc, "Текст cookie", T["cookieText"], "cookieText")
    add_plain(doc, f"Кнопки (только на новом сайте): {T['ckReject']} · {T['ckNeeded']} · {T['ckAll']}")

    add_h3(doc, "Часы работы (формулировка оригинала)")
    add_labeled_field(doc, "Часы", O["hours"], "hours")

    add_h3(doc, "Прочие UI-подписи")
    for key in (
        "more", "readOn", "allServices", "allCases", "allNews", "caseStudy",
        "scrollHint", "toTop", "showMore", "emptyNews", "emptyCases",
        "direction", "region", "outcome", "format", "projectWork", "regionValue",
        "callUs", "writeUs", "weekdays", "otherCases", "otherNews",
        "sources", "topics", "newer", "earlier", "related",
    ):
        add_plain(doc, f"{key}: {T[key]}")


def export_meta(doc):
    add_h2(doc, "SEO и meta-тексты страниц")
    pages = [
        ("Главная — title", C["homeTitle"]),
        ("Главная — description", C["homeDesc"], "homeDesc"),
        ("Главная — keywords", C["homeKw"]),
        ("Услуги — title", C["svcPageTitle"]),
        ("Услуги — description", C["svcPageDesc"]),
        ("О компании — title", f"О компании {O['name']} — эксперты по защите брендов с {FOUNDED} года"),
        ("Кейсы — title", f"Кейсы: борьба с контрафактом — {O['name']}"),
        ("Новости — title", f"Новости: борьба с контрафактом — {O['name']}"),
        ("Контакты — title", f"Контакты — {O['name']}, Москва"),
        ("Политика — title", C["privacyTitle"]),
        ("Политика — description", C["privacyDesc"]),
    ]
    for item in pages:
        label, text = item[0], item[1]
        rk = item[2] if len(item) > 2 else None
        add_labeled_field(doc, label, text, rk)


def export_home(doc):
    add_h2(doc, "Главная страница")
    add_kicker(doc, "Главный экран")
    add_h3(doc, C["heroTitle"].replace("\n", " "))
    add_plain(doc, C["heroPill"])
    add_labeled_field(doc, "Вступительный абзац (intro)", C["intro"], "intro")
    add_plain(doc, f"Подсказка прокрутки: {C['heroScroll']}")

    add_h3(doc, "Статистика")
    for s in SITE["stats"]:
        suffix = s.get("suffix", "") or s.get("unit", "")
        add_plain(doc, f"{s.get('n', '')}{suffix} — {s.get('label', '')}")

    add_h3(doc, "География")
    geo = SITE["geography"]
    add_kicker(doc, geo["kicker"])
    add_h4(doc, geo["title"])
    add_plain(doc, geo["lead"])
    add_plain(doc, "Страны: " + ", ".join(geo.get("regions", [])))
    for b in geo.get("bar", []):
        add_plain(doc, f"{b.get('n', '')} — {b.get('label', '')}")

    add_h3(doc, "Услуги (превью на главной)")
    add_kicker(doc, C["svcKicker"])
    add_labeled_field(doc, "Заголовок блока услуг", C["svcTitle"], "svcTitle")
    add_plain(doc, C["svcDesc"])
    add_plain(
        doc,
        "Примечание: вводный абзац «Единая методология…» на оригинале отсутствует — "
        "зелёная правка не предлагается.",
    )
    add_plain(doc, f"CTA: {T['allServices']}")

    add_h3(doc, "Подход и практика (превью)")
    add_kicker(doc, C["apprKicker"])
    add_h4(doc, C["apprTitle"])
    add_plain(doc, C["apprDesc"])
    add_plain(
        doc,
        "Примечание: блок «Подход и практика / отделы» — новый для сайта; на "
        "vlasta-s.com есть блок «Ключевые показатели за год» и тексты «О нас». "
        "Зелёные правки — в полном разделе ниже.",
    )

    add_h3(doc, "Кейсы (превью)")
    add_kicker(doc, C["casesKicker"])
    add_h4(doc, C["casesTitle"])
    add_plain(doc, f"CTA: {T['allCases']}")

    add_h3(doc, "История (превью)")
    add_kicker(doc, C["histKicker"])
    add_h4(doc, C["histTitle"])
    add_plain(doc, C["histHint"])

    add_h3(doc, "Новости (превью)")
    add_kicker(doc, C["newsKicker"])
    add_h4(doc, C["newsTitle"])
    add_plain(doc, f"CTA: {T['allNews']}")

    add_h3(doc, "Marquee клиентов")
    add_plain(doc, C["marquee"])


def export_services(doc):
    add_h2(doc, "Страница «Услуги» — полный текст")
    add_kicker(doc, C["svcKicker"])
    add_labeled_field(doc, "Заголовок раздела (как на главной)", C["svcTitle"], "svcTitle")
    add_h3(doc, C["svcH1"])
    add_plain(doc, C["svcLead"])
    add_plain(
        doc,
        "Примечание: H1/lead страницы услуг на оригинале нет в таком виде "
        "(там — «Наши услуги» + карточки направлений). Зелёная правка только "
        "для заголовка «Наши услуги» и ключевых оригинальных абзацев ниже.",
    )
    doc.add_paragraph()

    add_h4(doc, "Соответствие направлений оригиналу vlasta-s.com")
    add_ab_compare(
        doc,
        "Название направления ИС",
        ORIGINAL["serviceIp"],
        "Интеллектуальная собственность",
    )
    add_ab_compare(
        doc,
        "Корпоративная безопасность",
        ORIGINAL["serviceCorporate"],
        "Безопасность бизнеса (объединяет корпоративную и личную)",
    )
    add_ab_compare(
        doc,
        "Личная безопасность",
        ORIGINAL["servicePersonal"],
        "Включено в «Безопасность бизнеса» (отдельной плитки нет)",
    )
    add_ab_compare(
        doc,
        "Комплексная безопасность — заголовок",
        ORIGINAL["integratedSecurityTitle"],
        "Комплексная система безопасности",
        note=f"Альт. на оригинале: {ORIGINAL['integratedSecurityConcept']}",
    )
    doc.add_paragraph()

    for svc in SITE["services"]:
        add_h3(doc, f"{svc['num']}. {svc['title']} ({svc.get('block', '')})")

        if svc["id"] == "brand":
            add_labeled_field(
                doc,
                "Tagline / полный цикл защиты",
                strip_html(svc.get("tagline", "")),
                "brandTagline",
            )
            add_labeled_field(doc, "Intro", svc.get("intro", ""), "brandIntro")
        elif svc["id"] == "security":
            add_plain(doc, strip_html(svc.get("tagline", "")))
            add_labeled_field(doc, "Intro (корпоративная безопасность)", svc.get("intro", ""), "securityIntro")
            add_ab_compare(
                doc,
                "Личная безопасность — intro",
                ORIGINAL["personalSecurityIntro"],
                "",
                note="На новом сайте отдельного intro для личной безопасности нет; "
                "близкий текст — intro блока «Безопасность бизнеса» выше.",
            )
            add_ab_compare(
                doc,
                "Миссия «одного окна»",
                ORIGINAL["oneWindowMission"],
                svc.get("intro", ""),
                note="B — ближайший текущий intro блока «Безопасность бизнеса».",
            )
        elif svc["id"] == "intelligence":
            add_plain(doc, strip_html(svc.get("tagline", "")))
            add_labeled_field(doc, "Intro", svc.get("intro", ""), "intelligenceIntro")
        else:
            add_plain(doc, strip_html(svc.get("tagline", "")))
            add_plain(doc, svc.get("intro", ""))
            add_plain(
                doc,
                "Примечание: отдельного блока «Консалтинг» на оригинале нет "
                "(юридические услуги входят в ИС / корпоративную безопасность) — "
                "сравнение A/B не предлагается.",
            )

        highlights = svc.get("highlights", [])
        if highlights:
            add_h4(doc, "Ключевые пункты (плитки на главной)")
            add_bullets(doc, [strip_html(h.get("t", "")) for h in highlights])

        add_h4(doc, "Направления")
        for d in svc.get("directions", []):
            add_h4(doc, d["title"])
            if (
                d["title"] == "Комплексная система безопасности"
                and svc["id"] == "security"
            ):
                add_labeled_field(
                    doc,
                    "Текст направления",
                    d.get("text", ""),
                    "integratedSecurityText",
                )
                add_ab_compare(
                    doc,
                    "Заголовок направления",
                    ORIGINAL["integratedSecurityTitle"],
                    d["title"],
                )
            elif d["title"] == "Регистрация товарных знаков" and svc["id"] == "brand":
                add_labeled_field(
                    doc,
                    "Текст направления",
                    d.get("text", ""),
                    "trademarkRegistrationText",
                )
            else:
                add_plain(doc, d.get("text", ""))
            if d.get("bullets"):
                add_bullets(doc, d["bullets"])

        channels = svc.get("channels")
        if channels:
            add_h4(doc, channels.get("title", "Каналы мониторинга"))
            for item in channels.get("items", []):
                add_plain(doc, f"{item.get('h', '')}: {item.get('p', '')}")

        for flow in svc.get("flows", []):
            add_h4(doc, flow.get("title", ""))
            add_plain(doc, flow.get("note", ""))
            for step in flow.get("steps", []):
                add_plain(doc, f"• {step.get('h', '')}: {step.get('p', '')}")

        doc.add_paragraph()


def export_approach(doc):
    add_h2(doc, "Подход и практика — полный текст")
    add_kicker(doc, C["apprKicker"])
    add_h3(doc, C["apprTitle"])
    add_plain(doc, C["apprDesc"])
    add_plain(
        doc,
        "Примечание: заголовок «Все отделы — одна система» и описания отделов — "
        "новые для сайта; на vlasta-s.com прямых аналогов нет. Зелёная правка "
        "только для KPI-формулировок с оригинала.",
    )
    doc.add_paragraph()

    add_h4(doc, "Отделы (интерактивная схема) — без зелёной правки")
    for d in SITE["departments"]:
        add_h4(doc, f"{d.get('short', '')} — {d['title']}")
        add_plain(doc, d.get("text", ""))

    add_labeled_field(
        doc,
        "Заголовок блока показателей",
        "Показатели эффективности (KPI)",
        "apprKpiHeader",
    )
    kpi_current = "\n".join(
        f"{(r.get('prefix', '') or '')}{r.get('value', '')}{r.get('unit', '')} — "
        f"{strip_html(r.get('label', ''))}"
        for r in SITE.get("results", [])
    )
    add_ab_compare(
        doc,
        "KPI — подписи и метрики",
        ORIGINAL["kpiHeader"]
        + "\n• "
        + ORIGINAL["kpiDamage"]
        + "\n• "
        + ORIGINAL["kpiCases"]
        + "\n• "
        + ORIGINAL["kpiProductions"]
        + "\n• "
        + ORIGINAL["kpiBudget"],
        kpi_current,
        note="Набор метрик разный; правьте B, глядя на A, где смысл пересекается "
        "(ущерб / возврат затрат).",
    )

    add_h4(doc, "Этапы подхода")
    add_plain(
        doc,
        "На оригинале нет четырёхэтапной схемы «Анализ → Стратегия → "
        "Реализация → Сопровождение» — зелёная правка не предлагается.",
    )
    for step in SITE.get("approach", []):
        add_plain(
            doc,
            f"{step.get('n', '')}. {step.get('title', '')}: {step.get('text', '')}",
        )


def export_roadmap(doc):
    add_h2(doc, "История компании — дорожная карта")
    add_kicker(doc, C["histKicker"])
    add_h3(doc, C["histTitle"])
    add_plain(doc, C["histHint"])
    add_plain(
        doc,
        "Примечание: дорожная карта / timeline на vlasta-s.com в таком виде "
        "отсутствует — зелёная правка не предлагается.",
    )
    doc.add_paragraph()

    for item in SITE["timeline"]:
        add_h4(doc, f"{item['year']} — {item['title']}")
        add_plain(doc, item.get("text", ""))


def export_about(doc):
    add_h2(doc, "О компании — полный текст")
    add_kicker(doc, "О нас")
    add_h3(doc, C["aboutH1"])
    add_labeled_field(doc, "Lead (о компании)", C["aboutLead"], "aboutLead")
    doc.add_paragraph()

    add_kicker(doc, C["whoKicker"])
    add_h3(doc, C["whoTitle"])
    add_labeled_field(doc, "Кто мы", C["whoDesc"], "whoDesc")
    doc.add_paragraph()

    add_h4(doc, "Ценности")
    values_current = "\n".join(
        f"{v.get('num', '')}. {v.get('title', '')}: {v.get('text', '')}"
        for v in SITE.get("values", [])
    )
    add_ab_compare(
        doc,
        "Принципы / ценности",
        ORIGINAL["valuesPrinciple"],
        values_current,
        note="A — сводный принцип с оригинала; B — четыре отдельные ценности на новом сайте.",
    )
    add_kicker(doc, C["teamKicker"])
    add_h3(doc, C["teamTitle"])
    for m in SITE.get("team", []):
        add_h4(doc, m["name"])
        add_plain(doc, f"{m.get('role', '')}. {m.get('note', '')}")

    add_kicker(doc, C["assocKicker"])
    add_h3(doc, C["assocTitle"])
    add_plain(doc, C["assocDesc"])
    for a in SITE.get("associations", []):
        add_h4(doc, f"{a.get('abbr', '')} — {a.get('name', '')}")
        add_plain(doc, a.get("meta", ""))
        add_plain(doc, a.get("desc", ""))

    add_kicker(doc, C["clientsKicker"])
    add_h3(doc, C["clientsTitle"])
    add_plain(doc, strip_html(C["clientsNote"]))
    add_h4(doc, "Список клиентов (логотипы)")
    add_plain(doc, ", ".join(c["n"] for c in SITE.get("clients", [])))

    add_kicker(doc, C["lettersKicker"])
    add_h3(doc, C["lettersTitle"])
    add_plain(doc, C["lettersDesc"])
    add_h4(doc, "Благодарственные письма")
    for letter in SITE.get("letters", []):
        add_plain(doc, f"• {letter.get('name', '')}")


def export_contacts(doc):
    add_h2(doc, "Контакты")
    add_kicker(doc, C["contactsKick"])
    add_h3(doc, C["contactsH1"])
    add_plain(doc, C["contactsLead"])
    add_plain(
        doc,
        "Примечание: на vlasta-s.com страница контактов — адрес / тел. / e-mail "
        "без маркетингового H1 «Обсудим безопасность…». Зелёная правка H1/lead "
        "не предлагается. Слоган в подвале — см. раздел footer (миссия).",
    )
    add_plain(doc, f"{T['phone']}: {O['phone']}")
    add_plain(doc, f"{T['email']}: {O['email']}")
    add_ab_compare(doc, "Адрес", ORIGINAL["address"], O["address"])
    add_labeled_field(doc, "Часы", O["hours"], "hours")
    add_plain(doc, f"{T['callUs']} / {T['writeUs']}")


def export_privacy(doc):
    add_h2(doc, "Политика конфиденциальности")
    add_kicker(doc, C["privacyKick"])
    add_h3(doc, C["privacyH1"])
    add_plain(doc, C["privacyDesc"])
    doc.add_paragraph()
    for i, (heading, paras) in enumerate(PRIVACY, 1):
        add_h4(doc, f"{i:02d}. {heading}")
        for p in paras:
            add_plain(doc, p)


def export_cases(doc):
    add_h2(doc, "Кейсы — полный текст")
    add_kicker(doc, C["casesKicker"])
    add_h3(doc, C["casesH1"])
    add_plain(doc, C["casesLead"])
    doc.add_paragraph()

    for case in CASES:
        add_h3(doc, case.get("title", ""))
        add_plain(doc, f"Категория: {case.get('category', '')}")
        if case.get("metric"):
            add_plain(doc, f"Результат: {case.get('metric')}")
        if case.get("intro"):
            add_plain(doc, case["intro"])
        if case.get("outcome"):
            add_plain(doc, f"Итог: {case['outcome']}")
        if case.get("tags"):
            add_plain(doc, "Теги: " + ", ".join(case["tags"]))
        for sec in case.get("sections", []):
            add_h4(doc, sec.get("h", ""))
            for p in sec.get("p", []):
                add_plain(doc, p)
        doc.add_paragraph()


def export_news(doc):
    add_h2(doc, f"Новости — полный текст ({len(NEWS)} материалов)")
    add_kicker(doc, C["newsKicker"])
    add_h3(doc, C["newsH1"])
    add_plain(doc, C["newsLead"])
    doc.add_paragraph()

    for art in NEWS:
        add_h3(doc, art.get("title", ""))
        if art.get("h1") and art["h1"] != art.get("title"):
            add_plain(doc, f"H1: {art['h1']}")
        add_plain(doc, f"Дата: {art.get('dateDisp', '')} ({art.get('dateIso', '')})")
        add_plain(doc, f"Категория: {art.get('category', '')} · {art.get('cluster', '')}")
        if art.get("metaDesc"):
            add_plain(doc, f"Meta: {art['metaDesc']}")
        if art.get("excerpt"):
            add_plain(doc, f"Анонс: {art['excerpt']}")
        for block in art.get("body", []):
            if block.startswith("• "):
                doc.add_paragraph(block[2:], style="List Bullet")
            else:
                add_plain(doc, block)
        doc.add_paragraph()


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    export_preface(doc)
    add_section_break(doc)
    export_chrome(doc)
    add_section_break(doc)
    export_meta(doc)
    add_section_break(doc)
    export_home(doc)
    add_section_break(doc)
    export_services(doc)
    add_section_break(doc)
    export_approach(doc)
    add_section_break(doc)
    export_roadmap(doc)
    add_section_break(doc)
    export_about(doc)
    add_section_break(doc)
    export_contacts(doc)
    add_section_break(doc)
    export_privacy(doc)
    add_section_break(doc)
    export_cases(doc)
    add_section_break(doc)
    export_news(doc)

    out = ROOT / "docs/vlasta-site-content-review.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Saved: {out} ({len(NEWS)} news, {len(CASES)} cases)")


if __name__ == "__main__":
    main()
